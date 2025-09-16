import os
from docx import Document
from file_creator import FileCreator

class DocCreator(FileCreator):
    """Crea un archivo .docx en la carpeta Word."""

    def __init__(self):
        super().__init__("Word")

    def create_file(self, filename):
        doc = Document()
        doc.add_heading("Archivo de Prueba Word", level=1)
        doc.add_paragraph("Este es un archivo Word creado con python-docx.")
        file_path = os.path.join(self.folder, f"{filename}.docx")
        doc.save(file_path)
